# -*- coding: utf-8 -*-
"""Gerador de petições .docx no padrão visual do escritório Paulo R. Tercini Filho.

Replica a formatação de `base-peticao-previdenciaria-padrao-visual` com python-docx
(A4, Bookman Old Style 12, espaçamento 1,5, margens do escritório, timbre na 1ª
página, rodapé com endereço, títulos de seção em barra preta com texto branco,
recuo de primeira linha de 2 cm no judicial e 4 cm no CRPS).

Uso típico:
    from docx_escritorio import Peca
    p = Peca(recuo="jef")
    p.enderecamento("EXCELENTÍSSIMO ... JUIZADO ESPECIAL FEDERAL DE CATANDUVA ...")
    p.processo("Processo nº 0000000-00.0000.0.00.0000")
    p.paragrafo("FULANO DE TAL, ja qualificado ...")
    p.titulo("I - DA SINTESE")
    p.paragrafo("...")
    p.item("a) ...")
    p.data_local("Monte Alto, 23 de junho de 2026.")
    p.assinatura()
    p.salvar("/tmp/peca.docx")

Sem dois-pontos como separador logico no conteudo (regra do escritorio); isso e
responsabilidade de quem escreve o texto, o modulo so formata.
"""
from docx import Document
from docx.shared import Pt, Twips, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONTE = "Bookman Old Style"
RECUOS = {"jef": Twips(1134), "ordinario": Twips(1134), "ms": Twips(1134),
          "crps": Twips(2268), "nenhum": Twips(0)}


def _set_fonte(run, nome=FONTE, tamanho=12, negrito=False, cor=None):
    run.font.name = nome
    run.font.size = Pt(tamanho)
    run.bold = negrito
    if cor is not None:
        run.font.color.rgb = cor
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), nome)
    rfonts.set(qn("w:hAnsi"), nome)
    rfonts.set(qn("w:cs"), nome)


def _runs_md(paragraph, texto, nome=FONTE, tamanho=12, base_negrito=False, cor=None):
    """Adiciona o texto ao paragrafo interpretando **...** como negrito EMBUTIDO.
    Use o negrito so no que for realmente principal, para destacar. Fora dos
    marcadores, o texto segue `base_negrito`."""
    import re
    partes = re.split(r"\*\*(.+?)\*\*", texto)
    for i, parte in enumerate(partes):
        if parte == "":
            continue
        neg = base_negrito or (i % 2 == 1)  # trechos capturados (impares) = negrito
        _set_fonte(paragraph.add_run(parte), nome=nome, tamanho=tamanho,
                   negrito=neg, cor=cor)


class Peca:
    def __init__(self, recuo="jef"):
        doc = Document()
        sec = doc.sections[0]
        sec.page_width = Twips(11907)
        sec.page_height = Twips(16838)
        sec.top_margin = Twips(851)
        sec.bottom_margin = Twips(1134)
        sec.left_margin = Twips(1560)
        sec.right_margin = Twips(1134)
        sec.header_distance = Twips(720)
        sec.footer_distance = Twips(720)
        sec.different_first_page_header_footer = True

        normal = doc.styles["Normal"]
        normal.font.name = FONTE
        normal.font.size = Pt(12)
        rpr = normal.element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        for a in ("w:ascii", "w:hAnsi", "w:cs"):
            rfonts.set(qn(a), FONTE)
        pf = normal.paragraph_format
        pf.line_spacing = 1.5
        pf.space_before = Pt(12)
        pf.space_after = Pt(12)
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        self.doc = doc
        self.sec = sec
        self.recuo = RECUOS.get(recuo, RECUOS["jef"])
        self._timbre()
        self._rodape()

    def _timbre(self):
        hdr = self.sec.first_page_header
        linhas = [
            ("ADVOCACIA PREVIDENCIÁRIA", 14, True),
            ("PAULO ROBERTO TERCINI FILHO", 12, True),
            ("OAB/SP 331.110", 11, False),
        ]
        first = True
        for texto, tam, neg in linhas:
            p = hdr.paragraphs[0] if first else hdr.add_paragraph()
            first = False
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            _set_fonte(p.add_run(texto), tamanho=tam, negrito=neg)
        # borda inferior sob a ultima linha do timbre
        pPr = p._p.get_or_add_pPr()
        pbdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "6")
        bottom.set(qn("w:color"), "000000")
        pbdr.append(bottom)
        pPr.append(pbdr)

    def _rodape(self):
        ftr = self.sec.first_page_footer
        linhas = [
            "Rua Rui Barbosa, nº. 663, Centro, Monte Alto – SP",
            "Tel: 16-3242-2908 – Cel: 16-98140-9271",
        ]
        first = True
        for texto in linhas:
            p = ftr.paragraphs[0] if first else ftr.add_paragraph()
            first = False
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            _set_fonte(p.add_run(texto), tamanho=9)

    def enderecamento(self, texto):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _set_fonte(p.add_run(texto), negrito=True)
        self.doc.add_paragraph()
        return p

    def processo(self, texto, negrito=True):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _set_fonte(p.add_run(texto), negrito=negrito)
        self.doc.add_paragraph()
        return p

    def titulo(self, texto):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(8)
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "000000")
        pPr.append(shd)
        _set_fonte(p.add_run(" " + texto), tamanho=10, negrito=True,
                   cor=RGBColor(0xFF, 0xFF, 0xFF))
        return p

    def paragrafo(self, texto, recuar=True, negrito=False, centralizado=False):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if centralizado else WD_ALIGN_PARAGRAPH.JUSTIFY
        if recuar and not centralizado:
            p.paragraph_format.first_line_indent = self.recuo
        _runs_md(p, texto, base_negrito=negrito)
        return p

    def item(self, texto):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _runs_md(p, texto)
        return p

    def citacao(self, texto):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Twips(2268)
        r = p.add_run(texto)
        _set_fonte(r, tamanho=11)
        r.italic = True
        return p

    def data_local(self, texto):
        self.doc.add_paragraph()
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_fonte(p.add_run(texto))
        self.doc.add_paragraph()
        return p

    def assinatura(self):
        p1 = self.doc.add_paragraph()
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_fonte(p1.add_run("PAULO ROBERTO TERCINI FILHO"), negrito=True)
        p2 = self.doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_fonte(p2.add_run("OAB/SP 331.110"), negrito=True)

    def salvar(self, caminho):
        self.doc.save(caminho)
        return caminho
